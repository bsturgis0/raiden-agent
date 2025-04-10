from pathlib import Path
from typing import Optional
from langchain_core.tools import tool
import ffmpeg
import PyPDF2
from docx import Document
from PIL import Image
import pandas as pd

try:
    from app import WORKSPACE_DIR, color_text, _resolve_safe_path
except ImportError:
    WORKSPACE_DIR = Path("./raiden_workspace_srv")
    WORKSPACE_DIR.mkdir(exist_ok=True)
    def color_text(text, color="WHITE"): print(f"[{color}] {text}"); return text
    def _resolve_safe_path(filename: str) -> Path:
        target_path = (WORKSPACE_DIR / filename).resolve()
        if not str(target_path).startswith(str(WORKSPACE_DIR.resolve())):
            raise ValueError("Path traversal attempt detected.")
        return target_path

@tool
def convert_document(input_path: str, output_format: str) -> str:
    """Converts documents between formats (PDF, DOCX, TXT)."""
    try:
        input_path = _resolve_safe_path(input_path)
        if not input_path.exists():
            return f"Error: Input file {input_path} not found"

        output_filename = f"converted_{input_path.stem}.{output_format.lower()}"
        output_path = _resolve_safe_path(output_filename)
        
        input_ext = input_path.suffix.lower()[1:]
        output_format = output_format.lower()
        
        if input_ext == "pdf" and output_format == "docx":
            # PDF to DOCX
            doc = Document()
            with open(input_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    doc.add_paragraph(page.extract_text())
            doc.save(output_path)
            
        elif input_ext == "docx" and output_format == "pdf":
            # DOCX to PDF using docx2pdf
            from docx2pdf import convert
            convert(str(input_path), str(output_path))
            
        elif output_format == "txt":
            # Any to TXT
            if input_ext == "pdf":
                with open(input_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                    
            elif input_ext == "docx":
                doc = Document(input_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        
        return f"Successfully converted {input_path.name} to {output_filename}"
        
    except Exception as e:
        return f"Error converting document: {str(e)}"

@tool
def convert_image(input_path: str, output_format: str, quality: int = 95) -> str:
    """Converts images between formats with optional quality setting."""
    try:
        input_path = _resolve_safe_path(input_path)
        if not input_path.exists():
            return f"Error: Input file {input_path} not found"

        output_filename = f"converted_{input_path.stem}.{output_format.lower()}"
        output_path = _resolve_safe_path(output_filename)
        
        with Image.open(input_path) as img:
            # Convert RGBA to RGB if saving as JPEG
            if output_format.lower() == "jpeg" and img.mode == "RGBA":
                img = img.convert("RGB")
                
            img.save(output_path, quality=quality)
            
        return f"Successfully converted {input_path.name} to {output_filename}"
        
    except Exception as e:
        return f"Error converting image: {str(e)}"

@tool
def convert_data_format(input_path: str, output_format: str) -> str:
    """Converts data files between formats (CSV, JSON, XLSX, XML)."""
    try:
        input_path = _resolve_safe_path(input_path)
        if not input_path.exists():
            return f"Error: Input file {input_path} not found"

        output_filename = f"converted_{input_path.stem}.{output_format.lower()}"
        output_path = _resolve_safe_path(output_filename)
        
        # Read input
        if input_path.suffix.lower() == '.csv':
            df = pd.read_csv(input_path)
        elif input_path.suffix.lower() == '.json':
            df = pd.read_json(input_path)
        elif input_path.suffix.lower() == '.xlsx':
            df = pd.read_excel(input_path)
        elif input_path.suffix.lower() == '.xml':
            df = pd.read_xml(input_path)
        else:
            return f"Unsupported input format: {input_path.suffix}"
            
        # Write output
        if output_format.lower() == 'csv':
            df.to_csv(output_path, index=False)
        elif output_format.lower() == 'json':
            df.to_json(output_path, orient='records', indent=2)
        elif output_format.lower() == 'xlsx':
            df.to_excel(output_path, index=False)
        elif output_format.lower() == 'xml':
            df.to_xml(output_path)
        else:
            return f"Unsupported output format: {output_format}"
            
        return f"Successfully converted {input_path.name} to {output_filename}"
        
    except Exception as e:
        return f"Error converting data format: {str(e)}"
